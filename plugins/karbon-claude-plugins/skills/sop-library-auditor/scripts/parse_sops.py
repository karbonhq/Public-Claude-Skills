#!/usr/bin/env python3
"""
parse_sops.py — walk a folder of SOP files and emit a JSON inventory.

Usage:
    python parse_sops.py <folder> <output.json>

Supported formats:
    .md, .markdown, .txt   -- read directly as UTF-8
    .docx                  -- requires python-docx (pip install python-docx)
    .pdf                   -- requires pypdf (pip install pypdf)

Output JSON shape (one entry per file):
    {
      "inventory": [
        {
          "path": "absolute path",
          "relpath": "relative to root",
          "title": "best-guess title (first heading or filename)",
          "first_500_words": "...",
          "word_count": 1234,
          "last_modified": "2025-08-12T13:45:00",
          "ext": ".docx",
          "extraction_status": "ok" | "skipped" | "error: ...",
          "size_bytes": 12345
        },
        ...
      ],
      "summary": {
        "files_found": N,
        "files_extracted": N,
        "files_skipped": N,
        "files_errored": N
      }
    }

The skill reads this output and uses Claude to classify each entry into the
Domain A-J taxonomy. This script does not classify; it only extracts text.
"""

import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

SUPPORTED_TEXT_EXTS = {".md", ".markdown", ".txt", ".rst"}
SUPPORTED_DOCX_EXTS = {".docx"}
SUPPORTED_PDF_EXTS = {".pdf"}
SKIP_DIRS = {".git", "node_modules", "__pycache__", ".venv", "venv", "working", "interview-captures"}


def read_text_file(path: Path) -> str:
    """Read a UTF-8 text file, with a permissive fallback."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def read_docx(path: Path) -> str:
    """Extract text from a .docx file. Requires python-docx."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise RuntimeError("python-docx not installed. pip install python-docx")
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Capture table content too — many SOPs put RACI / inputs in tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    """Extract text from a .pdf file. Requires pypdf."""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        raise RuntimeError("pypdf not installed. pip install pypdf")
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def guess_title(text: str, fallback: str) -> str:
    """Pick the most likely document title."""
    if not text:
        return fallback
    # First-line markdown heading
    for line in text.splitlines()[:30]:
        line = line.strip()
        if line.startswith("# "):
            return line.lstrip("# ").strip()
        if line.startswith("## "):
            return line.lstrip("# ").strip()
    # First non-empty line under ~120 chars
    for line in text.splitlines()[:10]:
        line = line.strip()
        if 0 < len(line) < 120 and not line.startswith(("|", "---", "==")):
            return line
    return fallback


def first_n_words(text: str, n: int = 500) -> str:
    words = text.split()
    return " ".join(words[:n])


def extract_one(path: Path) -> dict:
    ext = path.suffix.lower()
    entry = {
        "path": str(path.resolve()),
        "title": path.stem,
        "first_500_words": "",
        "word_count": 0,
        "last_modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
        "ext": ext,
        "extraction_status": "ok",
        "size_bytes": path.stat().st_size,
    }
    try:
        if ext in SUPPORTED_TEXT_EXTS:
            text = read_text_file(path)
        elif ext in SUPPORTED_DOCX_EXTS:
            text = read_docx(path)
        elif ext in SUPPORTED_PDF_EXTS:
            text = read_pdf(path)
        else:
            entry["extraction_status"] = "skipped: unsupported format"
            return entry
        text = re.sub(r"\s+\n", "\n", text).strip()
        entry["title"] = guess_title(text, path.stem)
        entry["first_500_words"] = first_n_words(text, 500)
        entry["word_count"] = len(text.split())
    except Exception as exc:
        entry["extraction_status"] = f"error: {type(exc).__name__}: {exc}"
    return entry


def walk(root: Path):
    for dirpath, dirnames, filenames in os.walk(root):
        # Prune skip dirs in-place so os.walk doesn't descend
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS and not d.startswith(".")]
        for name in filenames:
            p = Path(dirpath) / name
            ext = p.suffix.lower()
            if ext in SUPPORTED_TEXT_EXTS | SUPPORTED_DOCX_EXTS | SUPPORTED_PDF_EXTS:
                yield p


def main():
    if len(sys.argv) != 3:
        print("Usage: python parse_sops.py <folder> <output.json>", file=sys.stderr)
        sys.exit(1)

    root = Path(sys.argv[1]).expanduser().resolve()
    out_path = Path(sys.argv[2]).expanduser().resolve()

    if not root.exists() or not root.is_dir():
        print(f"Folder not found or not a directory: {root}", file=sys.stderr)
        sys.exit(1)

    inventory = []
    for path in walk(root):
        entry = extract_one(path)
        try:
            entry["relpath"] = str(path.relative_to(root))
        except ValueError:
            entry["relpath"] = str(path)
        inventory.append(entry)

    summary = {
        "files_found": len(inventory),
        "files_extracted": sum(1 for e in inventory if e["extraction_status"] == "ok"),
        "files_skipped": sum(1 for e in inventory if e["extraction_status"].startswith("skipped")),
        "files_errored": sum(1 for e in inventory if e["extraction_status"].startswith("error")),
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8") as f:
        json.dump({"inventory": inventory, "summary": summary, "scanned_at": datetime.now().isoformat(timespec="seconds"), "root": str(root)}, f, indent=2)

    print(f"Scanned {summary['files_found']} files from {root}")
    print(f"  Extracted: {summary['files_extracted']}")
    print(f"  Skipped:   {summary['files_skipped']}")
    print(f"  Errored:   {summary['files_errored']}")
    print(f"Output: {out_path}")


if __name__ == "__main__":
    main()
